import streamlit as st
import os
import dotenv
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from time import time
import pandas as pd
from langchain.schema import Document
import logging
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import json



os.environ["USER_AGENT"] = "YourAppName/1.0"
# Load SpaCy model for semantic splitting

dotenv.load_dotenv()

DB_DOCS_LIMIT = 20  # Increase the document upload limit

from datetime import datetime

def get_current_date():
    """
    Returns the current date in a human-readable format.
    """
    return datetime.now().strftime("%B %d, %Y")



def extract_text_from_pdf(file_path):
    """
    Extract text from a PDF file, page by page.
    """
    try:
        with open(file_path, 'rb') as pdf:
            reader = PdfReader(pdf)
            return [page.extract_text().strip() for page in reader.pages]
    except Exception as e:
        logging.error(f"Failed to extract text from PDF: {e}")
        return []


    

def extract_text_from_docx(file_path):
    """Extract text from a Word document, including headings, paragraphs, and tables."""
    try:
        doc = DocxDocument(file_path)
        content = []
        current_section = None

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                if paragraph.style.name.startswith('Heading'):
                    if current_section and "content" in current_section and current_section["content"]:
                        content.append(current_section)
                    current_section = {"type": "heading", "text": text, "content": []}
                else:
                    if not current_section:
                        current_section = {"type": "text", "content": []}
                    if current_section["type"] == "text":
                        current_section["content"].append(text)
                    else:
                        content.append(current_section)
                        current_section = {"type": "text", "content": [text]}

        if current_section and "content" in current_section and current_section["content"]:
            content.append(current_section)

        return json.dumps(content, indent=4)
    except Exception as e:
        logging.error(f"Error processing Word document {file_path}: {e}")
        return json.dumps({"error": str(e)}, indent=4)


def clean_data(dataframe):
    """
    Clean the data in a Pandas DataFrame by:
    - Dropping rows/columns with excessive NaN values.
    - Filling NaN values with a placeholder.
    - Removing duplicate rows.
    - Resetting the index.
    """
    # Drop rows/columns with more than 50% NaN values
    dataframe = dataframe.dropna(thresh=dataframe.shape[1] * 0.5, axis=0)
    
    # Fill remaining NaN values with placeholders
    dataframe = dataframe.fillna("")
    
    # Drop duplicate rows
    dataframe = dataframe.drop_duplicates()
    
    # Reset the index
    dataframe.reset_index(drop=True, inplace=True)
    
    return dataframe

import chardet

def load_doc_to_db():
    """
    Handle document uploads. Use semantic chunking to split documents and store them in the vector database.
    """
    if "rag_docs" in st.session_state and st.session_state.rag_docs:
        docs = []  # List to store processed documents
        unique_sources = set(st.session_state.rag_sources)  # Track uploaded files to avoid duplicates

        for doc_file in st.session_state.rag_docs:
            if doc_file.name not in unique_sources and len(unique_sources) < DB_DOCS_LIMIT:
                os.makedirs("source_files", exist_ok=True)
                file_path = f"./source_files/{doc_file.name}"

                with open(file_path, "wb") as file:
                    file.write(doc_file.read())

                try:
                    if doc_file.type == "application/pdf":
                        raw_pages = extract_text_from_pdf(file_path)
                        if raw_pages.strip():
                            docs.append(Document(page_content=raw_pages))
                        else:
                            st.warning(f"No content extracted from {doc_file.name}.")

                    elif doc_file.name.endswith(".docx"):
                        raw_text = extract_text_from_docx(file_path)
                        if raw_text.strip():
                            docs.append(Document(page_content=raw_text))
                        else:
                            st.warning(f"No content extracted from {doc_file.name}.")

                    elif doc_file.type in ["text/plain", "text/markdown"]:
                        with open(file_path, "r", encoding="utf-8") as file:
                            raw_text = file.read()
                        docs.append(Document(page_content=raw_text))

                    elif doc_file.name.endswith((".xls", ".xlsx")):
                        excel_data = pd.read_excel(file_path, engine="openpyxl")
                        cleaned_data = clean_data(excel_data)
                        json_data = cleaned_data.to_json(orient="records")
                        docs.append(Document(page_content=json_data))

                    elif doc_file.name.endswith(".csv"):
                        try:
                            csv_data = pd.read_csv(file_path, encoding="utf-8")
                        except UnicodeDecodeError:
                            with open(file_path, 'rb') as f:
                                result = chardet.detect(f.read(10000))
                            csv_data = pd.read_csv(file_path, encoding=result['encoding'])

                        cleaned_data = clean_data(csv_data)
                        json_data = cleaned_data.to_json(orient="records")
                        docs.append(Document(page_content=json_data))

                    else:
                        st.warning(f"Document type {doc_file.type} not supported.")
                        continue

                    unique_sources.add(doc_file.name)
                    st.session_state.rag_sources = list(unique_sources)

                except Exception as e:
                    st.error(f"Error processing {doc_file.name}: {e}")

        if docs:
            # Use _split_and_load_docs to split documents and load into database
            _split_and_load_docs(docs, chunk_size=256, overlap_size=25)
            st.toast("Documents loaded successfully.", icon="✅")
        else:
            st.error(f"Maximum number of documents reached ({DB_DOCS_LIMIT}).")









def _split_and_load_docs(docs, chunk_size=256, overlap_size=25):
    """
    Split documents into sequential chunks based on fixed size with optional overlap and load into the vector database.
    """
    chunks = []

    for doc in docs:
        # Tokenize content into fixed-size chunks
        token_multiplier = 4  # Approximation: 4 characters ≈ 1 token
        max_char_count = chunk_size * token_multiplier
        overlap_char_count = overlap_size * token_multiplier

        text = doc.page_content
        start = 0

        while start < len(text):
            # End index for current chunk
            end = start + max_char_count
            chunks.append(text[start:end])

            # Move the start point with overlap
            start += max_char_count - overlap_char_count

    # Filter out empty chunks
    chunks = [chunk.strip() for chunk in chunks if chunk.strip()]

    # Store chunks in session state for display
    if "chunked_knowledge" not in st.session_state:
        st.session_state.chunked_knowledge = []
    st.session_state.chunked_knowledge.extend(chunks)

    # Create Document objects for each chunk
    document_chunks = [Document(page_content=chunk) for chunk in chunks]

    # Check if vector_db is initialized
    if "vector_db" not in st.session_state or st.session_state.vector_db is None:
        st.session_state.vector_db = initialize_vector_db(document_chunks)
    else:
        st.session_state.vector_db.add_documents(document_chunks)
        st.session_state.vector_db.persist()  # Persist changes to database





def initialize_vector_db(docs):
    """
    Initialize a vector database and store documents with embeddings and metadata.
    """
    # Add metadata to each document
    for i, doc in enumerate(docs):
        doc.metadata = {"index": i}

    # Ensure database folder exists
    os.makedirs("./chroma_db", exist_ok=True)

    # Generate embeddings
    embeddings = OpenAIEmbeddings(api_key=st.session_state.openai_api_key,
                                  model="text-embedding-3-small"
                                )
    document_texts = [doc.page_content for doc in docs]
    try:
        embeddings_data = embeddings.embed_documents(document_texts)
        if not embeddings_data:
            raise ValueError("Generated embeddings are empty.")
    except Exception as e:
        logging.error(f"Error generating embeddings: {e}")
        raise

    # Initialize Chroma vector database
    vector_db = Chroma.from_documents(
        documents=docs,
        embedding=embeddings,
        persist_directory="./chroma_db",  # Persist directory for Chroma
        collection_name=f"{str(time()).replace('.', '')[:14]}_{st.session_state['session_id']}",
    )
    vector_db.persist()  # Ensure changes are saved

    return vector_db





def _get_context_retriever_chain(vector_db, llm):
    retriever = vector_db.as_retriever(
        search_type="similarity",
        search_kwargs={"k": 3}  # Retrieve fewer documents for relevance, 
                                # Relevance: Increase k if you want to broaden the scope of retrieved documents.
                                # Efficiency: Decrease k if performance or relevance sufficiency is a concern.
    )

    prompt = ChatPromptTemplate.from_messages([
        MessagesPlaceholder(variable_name="messages"),
        ("user", "{input}"),
        ("system", "Use the retrieved knowledge to craft a relevant response."),
    ])

    retriever_chain = create_history_aware_retriever(llm, retriever, prompt)
    return retriever_chain



def get_conversational_rag_chain(llm):
    if "vector_db" not in st.session_state or not st.session_state.vector_db:
        raise ValueError("No vector database found. Please upload a knowledge base or use the default LLM.")

    retriever_chain = _get_context_retriever_chain(st.session_state.vector_db, llm)
    prompt = ChatPromptTemplate.from_messages([
        ("system", """
# Role
- You are a highly skilled and knowledgeable AI assistant named Zorro, designed to provide expert assistance to Axent employees across various departments. Your deep understanding of Axent's products, services, and internal processes allows you to deliver accurate and timely information to support the company's operations. But you also have the added flexability of being a standard claude model for people to interact with, only if their user query or question is not Axent related. Keep responses for these types of questions short and non-detailed, and do not tell them that their request seems unrelated to Axent's business, as they are using you for general use as well, not just Axent related topics.
- It's extremely important that you answer any question someone gives you about a general question as swiftly as possible, do not say but that "request seems unrelated to Axent's business", simply answer their question, even if it is unrelated to axent. 
- Your primary focus is to provide accurate answers without referencing where the data is sourced from.
- Under zero circumstances should you ever start the conversations stating "Based on [insert doucment]" or "according to [insert document]" when retrieving information from the knowledge base, the user should not know where you got the information from unless they specfiically asked you.
         

# Task
Assist Axent employees with their queries using the following step-by-step process:

1. Analyze the user's question to determine the type of information they are seeking.
2. Query the relevant Axent databases to retrieve the necessary data.
3. If the required information is found in the databases:
   a. For technical questions & any coding related questions, provide a thorough and accurate answer based on the PCB repair data and the PCB Repair Flowchart.
   b. For general questions, provide a brief and accurate response in 1-2 sentences.
4. If the information is not available in the databases, use your extensive knowledge as a Claude model to generate a helpful response based on the context provided.
5. Present the answer in a clear and professional manner, ensuring that it addresses the user's needs.
6. Ask the user if they have any additional questions.
7. If you are unable to retrieve any company-related queries, default back to a normal Claude model and answer as a regular LLM.

{context}
         
# Specifics
- Provide thorough and accurate responses to ensure the smooth operation of the company.
- Adhere to Axent's data privacy and security policies when accessing employee data.
- Prioritize timely and detailed responses for critical issues to minimize potential disruptions.
- Use the PCB Repair Flowchart to guide users through the troubleshooting process step-by-step, offering accurate and concise information. The general steps can be found from Example 1.
- For PCB repairs, include debugging recommendations and specific cases where others have had the same issue and how they fixed it.

# Context
- Axent is a premier electronic engineering company, specializing in the design, manufacture, installation and support of visual communication systems for various applications such as innovative road and safety digital signage.
- Axent specialises in designing and manufacturing innovative road and safety digital signage, including speed, bicycle, and wildlife awareness signs, school zone signs, fire danger rating signs, variable message signs and passenger information display, sporting scoreboards, service station price indicators and carpark displays.
- The company was founded in 1984 when the current director, Geoff Fontaine, built an electronic scoreboard display from his garage to automate the manual process of updating the scoreboard at the local circket centre that he worked at.
- As an AI assistant, your primary goal is to support Axent's employees by providing them with quick access to essential information, leveraging Axent's databases and your extensive knowledge base to streamline operations and enhance productivity.

# Examples
## Example 1 - PCB Repair Flow guidelines & Structure:
Q: How do I fix no heartbeat on a controller?
A: Based on the repair records, there are several common approaches to resolving a "No HB" (No Heartbeat) issue on Axent controllers:
1. Check the SD card slot and connection - sometimes a jump wire or resoldering the detect pin can resolve this issue.
2. Investigate Ram Chips - If the board fails to load the SD card, replacing both RAM chips often helps.
3. Examine the CPU - a faulty CPU can prevent heartbeat. Replacing the CPU and reflowing surrounding networking resistors has been successful in multiple cases.
4. Verify debug information - look at the debug port to understand why the PCB isn't booting properly.
Would you like me to elaborate on any of these troubleshooting steps? The specific solution can depend on the exact board model and symptoms.

## Example 2
Q: What is Axent?
A: Axent is a premier electronic engineering company, specializing in the design, manufacture, installation and support of visual communication systems for various applications such as innovative road and safety digital signage.

## Example 3
Q: Who is the founder of Axent?
A: The company was founded in 1984 when the current director, Geoff Fontaine, built an electronic scoreboard display from his garage to automate the manual process of updating the scoreboard at the local circket centre that he worked at.

# Notes
- Prioritize information retrieved from Axent's databases over general knowledge to ensure the most relevant and accurate answers.
- If unable to find the required information in the databases, use your Claude knowledge to provide a helpful response based on the available context.
- Maintain a professional and friendly tone in your interactions with Axent employees.
- Keep non-technical, generic questions concise, using only 1-2 sentences to provide accurate and relevant information.
- If users ask general questions, ensure your responses only mention what Axent does at an expected length of 1-2 sentences, and do not mention anything about the founder unless specifically asked.
- General Query Handling: If a user's query is unrelated to Axent's knowledge base or internal context, respond as a general-purpose AI (Claude model) with accurate and concise information. Keep responses as brief as possible (1-2 sentences) to align with the length and style of Axent-related responses. If the user requires additional details, follow up with a prompt like, "Would you like more information about this topic?" to encourage further dialogue.
- School signs are not radar activated they just automatically turn on at stipulated times of day.
- If the user asks a question about the date/what day it is, respond accurately with the current date from the function call, and **not the date from the documents unless specifically asked**, and **do not** mention that you are providing a simulated date based on the context of the conversation.
    """),  # Custom prompt, can modify for better compactness and token efficiency if needed.
        MessagesPlaceholder(variable_name="messages"),
        ("user", "{input}"),
    ])
    
    stuff_documents_chain = create_stuff_documents_chain(llm, prompt)

    return create_retrieval_chain(retriever_chain, stuff_documents_chain)


def extract_date_from_document(messages):
    """
    Extracts the document revision date from the processed document content.
    """
    for message in messages:
        if message.role == "assistant" and "document" in message.content.lower():
            # Search for dates in the document content
            import re
            date_pattern = r'\d{2}/\d{2}/\d{4}'  # Matches dates in DD/MM/YYYY format
            match = re.search(date_pattern, message.content)
            if match:
                return match.group(0)
    return None


def stream_llm_rag_response(llm_stream, messages):
    conversation_rag_chain = get_conversational_rag_chain(llm_stream)
    response_message = ""

    # Get the last user message
    user_message = messages[-1].content.lower()

    # Check if the user is asking for today's date
    if any(phrase in user_message for phrase in ["current date", "today's date", "date today"]):
        current_date = get_current_date()
        response_message = f"Today's date is {current_date}."
        yield response_message
        return

    # Check if the user is asking for a document-specific date
    if "document date" in user_message or "revision date" in user_message:
        # Fetch document-specific dates if available
        document_date = extract_date_from_document(messages)  # Custom function to extract document dates
        if document_date:
            response_message = f"The document's revision date is {document_date}."
        else:
            response_message = "I couldn't find a date in the document."
        yield response_message
        return

    try:
        # Use the RAG retriever chain
        for chunk in conversation_rag_chain.pick("answer").stream({"messages": messages[:-1], "input": messages[-1].content}):
            response_message += chunk
            yield chunk
    except ValueError:
        # Fallback to standard LLM response
        for chunk in llm_stream.stream(messages):
            response_message += chunk.content
            yield chunk

    return response_message
